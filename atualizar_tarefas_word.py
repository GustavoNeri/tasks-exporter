# -*- coding: utf-8 -*-
import os
import re
import sys
import math
import shutil
import datetime
from dotenv import load_dotenv

# Carregar variáveis de ambiente do arquivo .env
load_dotenv()

try:
    import urllib.parse as urlparse
except ImportError:
    import urlparse
import requests
from bs4 import BeautifulSoup
import openpyxl
from docx import Document
from docx.shared import Pt, RGBColor
from copy import copy
from copy import deepcopy

# ==========================================
# CONFIGURAÇÕES DO USUÁRIO
# ==========================================
USUARIO = os.getenv("TAREFAS_USER")
SENHA = os.getenv("TAREFAS_PASS")
IDS_FUNCIONARIOS = [65, 270, 271, 74, 75, 68]  # Adicione aqui os IDs dos funcionários (ex: [278, 279])
# 65  - Diovane Barbieri Gabriel
# 270 - Gustavo Neri
# 271 - Marlon David Domingos
# 74  - Sebastião Vitor dos Santos
# 75  - Thallys Henrique Lima da Silva
# 68  - Vítor Fontanari da Silva
# [65, 270, 271, 74, 75, 68]

# Mapa de ID para nome e cargo do funcionário
FUNCIONARIOS_MAP = {
    65:  {"nome": "Diovane Barbieri Gabriel", "cargo": "Analista de Implantação"},
    270: {"nome": "Gustavo Neri", "cargo": "Tech Lead e Coordenador"},
    271: {"nome": "Marlon David Domingos", "cargo": "Analista de Desenvolvimento"},
    74:  {"nome": "Sebastião Vitor dos Santos", "cargo": "Analista de Desenvolvimento"},
    75:  {"nome": "Thallys Henrique Lima da Silva", "cargo": "Analista de Desenvolvimento"},
    68:  {"nome": "Vítor Fontanari da Silva", "cargo": "Analista de Implantação"}
}
# ==========================================

def obter_periodo_proxima_semana():
    #hoje = datetime.datetime.strptime('2026-06-08', '%Y-%m-%d').date()
    hoje = datetime.date.today()
    # Próxima segunda-feira
    dias_para_segunda = 7 - hoje.weekday()
    proxima_segunda = hoje + datetime.timedelta(days=dias_para_segunda)
    # Próxima sexta-feira
    proxima_sexta = proxima_segunda + datetime.timedelta(days=4)
    return proxima_segunda, proxima_sexta

def encontrar_input(soup, selector, fallback_type=None, fallback_names=None):
    # Tenta usar o seletor exato
    elements = soup.select(selector)
    if elements:
        return elements[0]
        
    # Se falhar e contiver 'tbody', tenta sem o 'tbody' (parser do bs4 às vezes omite)
    if "tbody" in selector:
        selector_no_tbody = selector.replace("> tbody >", ">").replace("tbody >", "")
        elements = soup.select(selector_no_tbody)
        if elements:
            return elements[0]
            
    # Fallback por tipo de input
    if fallback_type:
        elements = soup.find_all("input", type=fallback_type)
        if elements:
            return elements[0]
            
    # Fallback por nome do atributo
    if fallback_names:
        for name in fallback_names:
            element = soup.find("input", attrs={"name": name})
            if element:
                return element
    return None

def normalizar_texto(texto):
    if not texto:
        return ""
    import unicodedata
    texto_str = str(texto).strip().rstrip(":").strip()
    nfkd_form = unicodedata.normalize('NFKD', texto_str)
    return "".join([c for c in nfkd_form if not unicodedata.combining(c)]).lower()

def encontrar_links_por_coluna(soup, nome_coluna="titulo"):
    links = []
    nome_coluna_norm = normalizar_texto(nome_coluna)
    for table in soup.find_all("table"):
        first_row = table.find("tr")
        if not first_row:
            continue
            
        # Pega apenas as células do cabeçalho (primeira linha da tabela)
        headers = [normalizar_texto(cell.get_text(strip=True)) for cell in first_row.find_all(["th", "td"])]
        
        col_idx = -1
        for idx, h in enumerate(headers):
            if nome_coluna_norm in h or h == nome_coluna_norm:
                col_idx = idx
                break
        
        if col_idx != -1:
            # Encontrou a coluna. Agora itera pelas linhas de dados (pulando o cabeçalho)
            for row in table.find_all("tr")[1:]:
                cells = row.find_all("td")
                if len(cells) > col_idx:
                    cell = cells[col_idx]
                    a = cell.find("a", href=True)
                    if a:
                        links.append(a["href"])
    return list(dict.fromkeys(links))

def extrair_links_tarefas(soup):
    task_links = []
    for a in soup.find_all("a", href=True):
        href = a["href"]
        # Padrões relaxados para corresponder a /atividade/ver, /tarefa/visualizar, etc.
        patterns = [
            "/atividade/ver", "/atividade/visualizar", "/atividade/detalhe", "/atividade/exibir", "id_atividade=",
            "/tarefa/ver", "/tarefa/visualizar", "/tarefa/detalhe", "/tarefa/exibir", "id_tarefa="
        ]
        if any(pat in href for pat in patterns):
            if "javascript:" not in href and "excluir" not in href and "deletar" not in href:
                task_links.append(href)
    return list(dict.fromkeys(task_links))

def extrair_texto_apos_dois_pontos(text):
    """Extrai apenas o texto após os dois pontos"""
    if ":" not in text:
        return text
    parts = text.split(":", 1)
    return parts[1].strip()

def obter_valor_select(soup, nome_campo):
    """
    Retorna o texto da opção selecionada de um SELECT.
    """

    select = soup.find("select", {"name": nome_campo})

    if not select:
        return ""

    for option in select.find_all("option"):
        if option.has_attr("selected"):
            return option.get_text(" ", strip=True)

    return ""

def extrair_dados_tarefa_por_seletores(soup):
    """Extrai dados da tarefa usando seletores corretos para a estrutura real"""
    dados = {}
    
    # Campos atuais dos selects
    status_atual = obter_valor_select(
        soup,
        "id_status"
    )
    if status_atual:
        dados["status"] = status_atual

    responsavel_atual = obter_valor_select(
        soup,
        "id_usuario_resp"
    )
    if responsavel_atual:
        dados["responsavel"] = responsavel_atual

    usuario_resp = obter_valor_select(
        soup,
        "id_usuario_para"
    )
    if usuario_resp:
        dados["incluido_por"] = usuario_resp

    prioridade = obter_valor_select(
        soup,
        "id_prioridade"
    )
    if prioridade:
        dados["prioridade"] = prioridade

    try:
        horas_input = soup.find("input", {"name": "horas_estimadas"})

        if horas_input:
            horas = horas_input.get("value", "").strip()

            if horas:
                dados["estimativa_raw"] = horas

    except Exception:
        pass

    try:
        # Empresa - da tabela barraTarefa, primeira coluna (td[0])
        empresa_elem = soup.select("table.barraTarefa tr td")
        if empresa_elem and len(empresa_elem) > 0:
            empresa_span = empresa_elem[0].find("span")
            if empresa_span:
                dados["empresa"] = empresa_span.get_text(strip=True)
    except Exception as e:
        pass
    
    try:
        # Projeto - da tabela barraTarefa, terceira coluna (td[2])
        projeto_elem = soup.select("table.barraTarefa tr td")
        if projeto_elem and len(projeto_elem) > 2:
            projeto_span = projeto_elem[2].find("span")
            if projeto_span:
                dados["projeto"] = projeto_span.get_text(strip=True)
    except Exception as e:
        pass
    
    try:
        # Atividade/Titulo - procura por <strong> que vem após "Título:"
        titulo_div = None
        for div in soup.find_all("div", class_="titulo"):
            if "Título" in div.get_text():
                titulo_div = div
                break
        
        if titulo_div:
            strong_elem = titulo_div.find_next("strong")
            if strong_elem:
                dados["atividade"] = strong_elem.get_text(strip=True)
        else:
            titulo_elem = soup.find("strong", string=lambda x: x and len(x) > 20 and "Empresa" not in x)
            if titulo_elem:
                dados["atividade"] = titulo_elem.get_text(strip=True)
    except Exception as e:
        pass
    
    try:
        # Bloco de informacoes - procura por ul class="info-container"
        info_ul = soup.find("ul", class_="info-container")
        
        if info_ul:
            lis = info_ul.find_all("li")
            
            if len(lis) > 2:
                incluido_em_text = lis[2].get_text(strip=True)
                match_data = re.search(r"\d{1,2}/\d{1,2}/\d{4}", incluido_em_text)
                if match_data:
                    dados["incluido_em"] = match_data.group(0)
            
            if len(lis) > 7:
                entrega_em_text = lis[7].get_text(strip=True)
                match_data = re.search(r"\d{1,2}/\d{1,2}/\d{4}", entrega_em_text)
                if match_data:
                    dados["entrega_em"] = match_data.group(0)
            
            if len(lis) > 8:
                ordem_text = lis[8].get_text(strip=True)
                dados["ordem"] = extrair_texto_apos_dois_pontos(ordem_text)

    except Exception as e:
        print("  ! Erro ao extrair informacoes: " + str(e))
        import traceback
        traceback.print_exc()
    
    # Extrair descrição da tarefa
    try:
        descricao_tarefa = extrair_descricao_tarefa(soup)
        if descricao_tarefa:
            dados["descricao_tarefa"] = descricao_tarefa
        else:
            dados["descricao_tarefa"] = ""
    except Exception as e:
        dados["descricao_tarefa"] = ""
    
    return dados

def extrair_descricao_tarefa(soup):
    """Extrai a descrição da tarefa usando o seletor CSS específico"""
    try:
        # Seletor: body > div.divPrincipal > div.divConteudo > form > fieldset > table:nth-child(3) > tbody > tr:nth-child(3) > td > div.descricao_tarefa
        descricao_elem = soup.select("div.descricao_tarefa")
        if descricao_elem:
            descricao = descricao_elem[0].get_text().strip()
            return descricao
        
        # Fallback: procura por div com class contendo "descricao" ou "description"
        for div in soup.find_all("div"):
            classes = div.get("class", [])
            if isinstance(classes, list):
                classes_str = " ".join(classes)
            else:
                classes_str = str(classes)
            
            if "descricao" in classes_str.lower() or "description" in classes_str.lower():
                text = div.get_text().strip()
                if text and len(text) > 10:
                    return text
        
        return ""
    except Exception as e:
        print("  ! Erro ao extrair descricao: " + str(e))
        return ""

def formatar_run(run, font_name="Verdana", size_pt=10, bold=False, italic=False, color_rgb=None):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color_rgb:
        run.font.color.rgb = color_rgb

def editar_docx_funcionario(template_path, output_path, funcionario_id, lista_tarefas, segunda_file_str, sexta_file_str):
    """Edita o template DOCX com dados de múltiplas tarefas de um funcionário"""
    try:
        from copy import deepcopy
        
        if not lista_tarefas:
            return False
        
        # Carregar documento
        doc = Document(template_path)
        
        # Obter dados do funcionário
        func_info = FUNCIONARIOS_MAP.get(funcionario_id, {"nome": "Funcionário", "cargo": "Cargo"})
        func_nome = func_info["nome"]
        func_cargo = func_info["cargo"]
        
        # 1. Substituir cabeçalho com formatação específica
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if "Função / Cargo do prestador" in text or "Cargo do prestador" in text or "TAREFAS DA SEMANA" in text:
                paragraph.text = ""
                r1 = paragraph.add_run("TAREFAS DA SEMANA")
                formatar_run(r1, font_name="Verdana", size_pt=14, bold=True)
                r2 = paragraph.add_run(f" – {func_cargo} - {func_nome.upper()}")
                formatar_run(r2, font_name="Verdana", size_pt=12, bold=True)
            elif "Período: XX-XX-XXXX a XX-XX-XXXX" in text or "Período:" in text:
                paragraph.text = ""
                r = paragraph.add_run("Período: " + segunda_file_str + " a " + sexta_file_str)
                formatar_run(r, font_name="Verdana", size_pt=12, bold=True)
            elif text == "NOME DO PRESTADOR" or text == func_nome.upper():
                paragraph.text = ""
                r = paragraph.add_run(func_nome.upper())
                formatar_run(r, font_name="Verdana", size_pt=12, bold=True)
        
        # 2. Encontrar índices de bloco e Total
        bloco_start_idx = -1
        bloco_end_idx = -1
        total_idx = -1
        
        for i, p in enumerate(doc.paragraphs):
            text = p.text.strip()
            if "Categoria-Subcategoria" in text and "Descrição/Título" in text:
                bloco_start_idx = i
            elif bloco_start_idx != -1 and bloco_end_idx == -1 and "Estimativa:" in text and "XX h" in text:
                bloco_end_idx = i
            if "Total: XX horas" in text:
                total_idx = i
        
        if bloco_start_idx == -1 or bloco_end_idx == -1 or total_idx == -1:
            return False
        
        # 3. Duplicar o bloco template para as tarefas extras (tarefa 1 em diante)
        # IMPORTANTE: Fazemos isso ANTES de modificar o bloco original, para que as cópias
        # sejam idênticas ao template e contenham os placeholders originais.
        total_elem = doc.paragraphs[total_idx]._element
        for idx_tarefa in range(1, len(lista_tarefas)):
            for i in range(bloco_start_idx, bloco_end_idx + 1):
                elem_original = doc.paragraphs[i]._element
                elem_copia = deepcopy(elem_original)
                total_elem.addprevious(elem_copia)
                
        # 4. Preencher as tarefas
        # Agora que todos os blocos foram criados na estrutura do XML, recarregamos
        # doc.paragraphs como uma lista estática para iterar e substituir os valores corretos.
        paragraphs_to_delete = []
        tarefa_idx = 0
        
        i = 0
        paragraphs = list(doc.paragraphs)
        num_paragraphs = len(paragraphs)
        
        while i < num_paragraphs:
            p = paragraphs[i]
            text = p.text.strip()
            if "Categoria-Subcategoria" in text and "Descrição/Título" in text:
                if tarefa_idx < len(lista_tarefas):
                    tarefa = lista_tarefas[tarefa_idx]
                    
                    # Categoria-Subcategoria > Descrição/Título da atividade
                    p.text = ""
                    r = p.add_run(tarefa.get("atividade", ""))
                    formatar_run(r, font_name="Verdana", size_pt=10, bold=True)
                    
                    # Descrição da tarefa (rótulo)
                    p_desc_label = paragraphs[i + 1]
                    p_desc_label.text = ""
                    r_lbl = p_desc_label.add_run("Descrição da tarefa:")
                    formatar_run(r_lbl, font_name="Verdana", size_pt=10, bold=False)
                    
                    # Descrição da tarefa (conteúdo)
                    # O campo da descrição no site é mapeado para tarefa["descricao_tarefa"]
                    desc_raw = tarefa.get("descricao_tarefa", "")
                    
                    # Split por novas linhas e strip de espaços em branco
                    desc_lines = [line.strip() for line in desc_raw.split('\n')]
                    while desc_lines and not desc_lines[-1]:
                        desc_lines.pop()
                    while desc_lines and not desc_lines[0]:
                        desc_lines.pop(0)
                    if not desc_lines:
                        desc_lines = [""]
                    
                    # Temos exatamente 5 parágrafos de placeholder de descrição:
                    # p[i+2], p[i+3], p[i+4], p[i+5], p[i+6]
                    p_placeholders = [
                        paragraphs[i + 2],
                        paragraphs[i + 3],
                        paragraphs[i + 4],
                        paragraphs[i + 5],
                        paragraphs[i + 6]
                    ]
                    
                    if len(desc_lines) <= 5:
                        for j in range(5):
                            if j < len(desc_lines):
                                p_placeholders[j].text = ""
                                r = p_placeholders[j].add_run(desc_lines[j])
                                formatar_run(r, font_name="Verdana", size_pt=10, bold=False)
                            else:
                                paragraphs_to_delete.append(p_placeholders[j])
                    else:
                        for j in range(4):
                            p_placeholders[j].text = ""
                            r = p_placeholders[j].add_run(desc_lines[j])
                            formatar_run(r, font_name="Verdana", size_pt=10, bold=False)
                        
                        p_placeholders[4].text = ""
                        r = p_placeholders[4].add_run("\n".join(desc_lines[4:]))
                        formatar_run(r, font_name="Verdana", size_pt=10, bold=False)
                    
                    # Estimativa
                    p_estimativa = paragraphs[i + 8]
                    est = tarefa.get("estimativa_horas", 0)
                    if isinstance(est, float) and not est.is_integer():
                        est_str = str(est)
                    else:
                        est_str = str(int(est))
                    
                    p_estimativa.text = ""
                    r_est = p_estimativa.add_run("Estimativa: " + est_str + " h")
                    formatar_run(r_est, font_name="Verdana", size_pt=10, bold=False)
                    
                    tarefa_idx += 1
                
                # Pula o tamanho do bloco para otimizar a busca
                i += 9
            else:
                i += 1
                
        # 5. Remover os parágrafos de placeholder que não foram utilizados
        for p_del in paragraphs_to_delete:
            try:
                p_del._element.getparent().remove(p_del._element)
            except Exception as e:
                # Caso já tenha sido removido de alguma forma
                pass
        
        # 6. Atualizar Total
        total_horas = sum(t.get("estimativa_horas", 0) for t in lista_tarefas)
        if isinstance(total_horas, float) and total_horas.is_integer():
            total_str = str(int(total_horas))
        else:
            total_str = str(total_horas)
        
        # Encontrar Total novamente (sua posição mudou devido a duplicações/remoções)
        for p in doc.paragraphs:
            if "Total:" in p.text and ("XX horas" in p.text or "horas" in p.text):
                p.text = ""
                r = p.add_run("Total: " + total_str + " horas")
                formatar_run(r, font_name="Verdana", size_pt=10, bold=True)
                break
        
        # 7. Salvar
        doc.save(output_path)
        return True
    except Exception as e:
        print("  ! Erro ao editar DOCX: " + str(e))
        import traceback
        traceback.print_exc()
        return False

def main():
    if not USUARIO or not SENHA:
        print("Erro: por favor configure as variáveis de ambiente TAREFAS_USER e TAREFAS_PASS.")
        sys.exit(1)
        
    # Calcular datas do Word (mantem a regra original da proxima semana)
    segunda, sexta = obter_periodo_proxima_semana()
    periodo_de_str = segunda.strftime("%d/%m/%Y")
    periodo_ate_str = sexta.strftime("%d/%m/%Y")
    
    segunda_file_str = segunda.strftime("%d-%m-%Y")
    sexta_file_str = sexta.strftime("%d-%m-%Y")
    
    print("Processando Word semana de " + periodo_de_str + " a " + periodo_ate_str)
    
    # Verifica diretório de trabalho (deve ser o mesmo do script para acessar os arquivos locais)
    workspace_dir = os.path.dirname(os.path.abspath(__file__))
    
    # Iniciar sessão
    session = requests.Session()
    session.headers.update({
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/100.0.0.0 Safari/537.36"
    })
    
    # 1. Login disparado pelo redirect de acesso à lista
    primeiro_id = IDS_FUNCIONARIOS[0] if IDS_FUNCIONARIOS else 278
    list_redirect_url = "https://tasks.amlgroup.com.br/atividade/lista?id_usuario_responsavel=" + str(primeiro_id) + "&periodoDe=" + periodo_de_str + "&periodoAte=" + periodo_ate_str
    try:
        response = session.get(list_redirect_url)
    except Exception as e:
        print("ERRO ao conectar ao portal: " + str(e))
        sys.exit(1)
        
    soup = BeautifulSoup(response.text, "html.parser")
    form = soup.find("form", id="form1") or soup.find("form")
    
    form_data = {}
    if form:
        for ipt in form.find_all("input"):
            name = ipt.get("name")
            if name:
                form_data[name] = ipt.get("value", "")
        for btn in form.find_all("input", type="submit"):
            name = btn.get("name")
            if name:
                form_data[name] = btn.get("value", "")
                
    user_selector = "#form1 > table > tbody > tr:nth-child(2) > td > table > tbody > tr:nth-child(1) > td:nth-child(2) > input"
    pass_selector = "#form1 > table > tbody > tr:nth-child(2) > td > table > tbody > tr:nth-child(2) > td:nth-child(2) > input"
    
    user_input = encontrar_input(soup, user_selector, fallback_type="text", fallback_names=["usuario", "login", "email"])
    pass_input = encontrar_input(soup, pass_selector, fallback_type="password", fallback_names=["senha", "password"])
    
    if user_input:
        user_name_attr = user_input.get("name")
        if user_name_attr:
            form_data[user_name_attr] = USUARIO
    if pass_input:
        pass_name_attr = pass_input.get("name")
        if pass_name_attr:
            form_data[pass_name_attr] = SENHA
            
    action_url = form.get("action", "") if form else ""
    submit_url = urlparse.urljoin(response.url, action_url) if action_url else response.url
    
    try:
        login_resp = session.post(submit_url, data=form_data)
    except Exception as e:
        print("ERRO ao fazer login: " + str(e))
        sys.exit(1)
        
    if "senha" in login_resp.text.lower() and "login" in login_resp.text.lower() and login_resp.url == response.url:
        print("ERRO: Autenticacao falhou")
        sys.exit(1)
    
    total_inserido = 0
    tarefas_por_funcionario = {}  # { emp_id: [lista de tarefas] }
    
    for emp_id in IDS_FUNCIONARIOS:
        tarefas_por_funcionario[emp_id] = []      
            
    # Buscar tarefas do Word usando o periodo semanal original
    for emp_id in IDS_FUNCIONARIOS:
        func_info = FUNCIONARIOS_MAP.get(emp_id, {"nome": "Funcionario"})
        print("Extraindo Word: " + func_info["nome"])
        list_url = "https://tasks.amlgroup.com.br/atividade/lista"
        params = {
            "id_empresa_grupo": "",
            "id_empresa_area": "",
            "id_usuario_responsavel": str(emp_id),
            "id_projeto": "",
            "status": "",
            "periodoDe": periodo_de_str,
            "periodoAte": periodo_ate_str,
            "orderby": "tarefa_informacao.data_entrega, tarefa_informacao.ordem ASC"
        }
        
        try:
            list_resp = session.get(list_url, params=params)
        except Exception as e:
            print("  ERRO ao buscar tarefas do Word: " + str(e))
            continue
            
        if list_resp.status_code != 200:
            print("  ERRO Word: status " + str(list_resp.status_code))
            continue
            
        list_soup = BeautifulSoup(list_resp.text, "html.parser")
        task_links = encontrar_links_por_coluna(list_soup, "titulo")
        if not task_links:
            task_links = encontrar_links_por_coluna(list_soup, "título")
        if not task_links:
            task_links = encontrar_links_por_coluna(list_soup, "atividade")
        if not task_links:
            task_links = extrair_links_tarefas(list_soup)
            
        for link in task_links:
            task_url = urlparse.urljoin(list_resp.url, link)
            try:
                task_resp = session.get(task_url)
            except Exception as e:
                print("  ERRO ao carregar tarefa do Word: " + str(e))
                continue
                
            if task_resp.status_code != 200:
                continue
                
            task_soup = BeautifulSoup(task_resp.text, "html.parser")
            dados_tarefa = extrair_dados_tarefa_por_seletores(task_soup)
            estimativa_raw = dados_tarefa.get("estimativa_raw", "0")
            horas = float(estimativa_raw) if estimativa_raw else 0.0
            estimativa = int(horas) if horas.is_integer() else horas
            
            tarefas_por_funcionario[emp_id].append({
                "atividade": dados_tarefa.get("atividade", ""),
                "descricao_tarefa": dados_tarefa.get("descricao_tarefa", ""),
                "estimativa_horas": estimativa
            })
            
    # Agora gerar DOCX por funcionário (com múltiplas tarefas)
    
    docx_template_name = "00-Tarefas-Semana-XX-XX-XXXX-a-XX-XX-XXXX–Nome-do-Prestador.docx"
    docx_template_path = os.path.join(workspace_dir, docx_template_name)
    
    for emp_id, lista_tarefas in tarefas_por_funcionario.items():
        if not lista_tarefas:
            continue
        
        func_info = FUNCIONARIOS_MAP.get(emp_id, {"nome": "Funcionario"})
        func_nome_arquivo = func_info["nome"].replace(" ", "-")
        
        docx_nome = "Tarefas-Semana-" + segunda_file_str + "-a-" + sexta_file_str + "-" + func_nome_arquivo + ".docx"
        docx_output_path = os.path.join(workspace_dir, docx_nome)
        
        if os.path.exists(docx_template_path):
            if editar_docx_funcionario(docx_template_path, docx_output_path, emp_id, lista_tarefas, segunda_file_str, sexta_file_str):
                print("Word: " + func_info["nome"])
            else:
                print("  ERRO ao gerar Word para " + func_info["nome"])
        else:
            print("  ERRO: Template Word nao encontrado")
    
    # Limpar arquivo de debug ao finalizar o processo
    debug_file = "debug_page.html"
    if os.path.exists(debug_file):
        try:
            os.remove(debug_file)
        except Exception as e:
            print("ERRO ao remover arquivo de debug: " + str(e))
    
    print("\nProcessamento concluido!")

if __name__ == "__main__":
    main()
